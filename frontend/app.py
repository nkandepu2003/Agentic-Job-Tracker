# app.py
# AI Agentic Job Tracker Dashboard

import streamlit as st
import pandas as pd
import sys
import os
import re
import io

sys.path.append(os.path.dirname(
    os.path.dirname(os.path.abspath(__file__))))

from database.models import get_db, JobApplication
from datetime import datetime

# ─────────────────────────────────────
# PAGE SETTINGS
# ─────────────────────────────────────
st.set_page_config(
    page_title="AI Job Tracker",
    page_icon="💼",
    layout="wide"
)

# ─────────────────────────────────────
# TITLE
# ─────────────────────────────────────
st.title("💼 AI Agentic Job Tracker")
st.markdown(
    "Track every application. Never miss a follow up.")

# ─────────────────────────────────────
# SIDEBAR
# ─────────────────────────────────────
st.sidebar.title("➕ Add New Application")

company = st.sidebar.text_input("Company Name")
job_title = st.sidebar.text_input("Job Title")
job_url = st.sidebar.text_input("Job URL (optional)")
location = st.sidebar.text_input("Location (optional)")
salary = st.sidebar.text_input("Salary Range (optional)")
notes = st.sidebar.text_area("Notes (optional)")

status = st.sidebar.selectbox(
    "Status",
    ["Applied", "Interview Scheduled",
     "Interview Done", "Offer Received",
     "Rejected", "Withdrawn"]
)

if st.sidebar.button("💾 Save Application"):
    if company and job_title:
        db = get_db()
        new_application = JobApplication(
            company=company,
            job_title=job_title,
            job_url=job_url,
            location=location,
            salary=salary,
            notes=notes,
            status=status,
            date_applied=datetime.now()
        )
        db.add(new_application)
        db.commit()
        db.close()
        st.sidebar.success("✅ Application saved!")
    else:
        st.sidebar.error(
            "❌ Please fill Company and Job Title!")

# ─────────────────────────────────────
# MAIN AREA
# ─────────────────────────────────────
db = get_db()
applications = db.query(JobApplication).all()
db.close()

# ─────────────────────────────────────
# STATS ROW
# ─────────────────────────────────────
st.subheader("📊 Your Job Hunt Stats")

col1, col2, col3, col4 = st.columns(4)

total = len(applications)
interviews = len([a for a in applications
                  if "Interview" in a.status])
offers = len([a for a in applications
              if a.status == "Offer Received"])
rejected = len([a for a in applications
                if a.status == "Rejected"])

with col1:
    st.metric("📨 Total Applied", total)
with col2:
    st.metric("🎯 Interviews", interviews)
with col3:
    st.metric("🎉 Offers", offers)
with col4:
    st.metric("❌ Rejected", rejected)

st.divider()

# ─────────────────────────────────────
# APPLICATIONS TABLE
# ─────────────────────────────────────
st.subheader("📋 All Applications")

if applications:
    data = []
    for app in applications:
        data.append({
            "Company": app.company,
            "Job Title": app.job_title,
            "Status": app.status,
            "Location": app.location or "N/A",
            "Salary": app.salary or "N/A",
            "Date Applied": app.date_applied.strftime(
                "%d %b %Y"
            ) if app.date_applied else "N/A",
            "Notes": app.notes or "N/A"
        })
    df = pd.DataFrame(data)
    st.dataframe(df, use_container_width=True)
else:
    st.info(
        "👆 Add your first job application "
        "using the sidebar!")

st.divider()

# ─────────────────────────────────────
# STATUS FILTER
# ─────────────────────────────────────
st.subheader("🔍 Filter by Status")

selected_status = st.selectbox(
    "Show applications with status:",
    ["All", "Applied", "Interview Scheduled",
     "Interview Done", "Offer Received",
     "Rejected", "Withdrawn"]
)

if selected_status != "All" and applications:
    filtered = [a for a in applications
                if a.status == selected_status]
    if filtered:
        data_filtered = []
        for app in filtered:
            data_filtered.append({
                "Company": app.company,
                "Job Title": app.job_title,
                "Status": app.status,
                "Date Applied": app.date_applied.strftime(
                    "%d %b %Y"
                ) if app.date_applied else "N/A",
                "Notes": app.notes or "N/A"
            })
        st.dataframe(
            pd.DataFrame(data_filtered),
            use_container_width=True
        )
    else:
        st.info(
            f"No applications with status: "
            f"{selected_status}")

st.divider()

# ─────────────────────────────────────
# RESUME TAILOR AGENT
# ─────────────────────────────────────
st.subheader("🤖 AI Resume Tailor Agent")
st.markdown(
    "Paste a job description and AI will "
    "tailor your resume instantly!")

job_description = st.text_area(
    "Paste Job Description Here:",
    height=200,
    placeholder="Copy and paste the full job description here...",
    key="tailor_jd"
)

if st.button("✨ Tailor My Resume with AI"):
    if job_description:
        with st.spinner(
            "🤖 AI is tailoring your resume..."
        ):
            try:
                from agents.tailor_agent import (
                    run_tailor_agent)
                resume_path = "resume.pdf"
                tailored_resume, error = run_tailor_agent(
                    resume_path, job_description)

                if error:
                    st.error(f"❌ Error: {error}")
                else:
                    st.success(
                        "✅ Your tailored resume is ready!")
                    st.subheader("👀 Preview (Formatted)")
                    st.markdown(tailored_resume)
                    st.divider()

                    col1, col2 = st.columns(2)
                    with col1:
                        st.download_button(
                            label="⬇️ Download as TXT",
                            data=tailored_resume,
                            file_name="tailored_resume.txt",
                            mime="text/plain"
                        )
                    with col2:
                        if st.button(
                            "📄 Download as Word Doc",
                            key="tailor_word"
                        ):
                            try:
                                from docx import Document
                                doc = Document()
                                for line in tailored_resume.split('\n'):
                                    if not line.strip():
                                        doc.add_paragraph('')
                                        continue
                                    if (line.strip().startswith('*') and
                                            not line.strip().startswith('**')):
                                        p = doc.add_paragraph(
                                            style='List Bullet')
                                        line = line.strip().lstrip(
                                            '*').strip()
                                    else:
                                        p = doc.add_paragraph()
                                    parts = re.split(
                                        r'\*\*(.*?)\*\*', line)
                                    for i, part in enumerate(parts):
                                        if i % 2 == 0:
                                            p.add_run(part)
                                        else:
                                            run = p.add_run(part)
                                            run.bold = True
                                doc_bytes = io.BytesIO()
                                doc.save(doc_bytes)
                                doc_bytes.seek(0)
                                st.download_button(
                                    label="⬇️ Save Word File",
                                    data=doc_bytes.getvalue(),
                                    file_name="tailored_resume.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key="tailor_docx"
                                )
                            except Exception as e:
                                st.error(f"❌ Error: {str(e)}")
            except Exception as e:
                st.error(
                    f"❌ Something went wrong: {str(e)}")
    else:
        st.warning(
            "⚠️ Please paste a job description first!")

st.divider()

# ─────────────────────────────────────
# COVER LETTER AGENT
# ─────────────────────────────────────
st.subheader("✉️ AI Cover Letter Agent")
st.markdown(
    "Enter company name and job description — "
    "AI writes your cover letter!")

cover_company = st.text_input(
    "Company Name:",
    placeholder="e.g. Google, Amazon, Microsoft...",
    key="cover_company"
)

cover_jd = st.text_area(
    "Paste Job Description Here:",
    height=200,
    placeholder="Copy and paste the full job description here...",
    key="cover_jd"
)

if st.button("✉️ Write My Cover Letter"):
    if cover_company and cover_jd:
        with st.spinner(
            "✍️ AI is writing your cover letter..."
        ):
            try:
                from agents.cover_letter_agent import (
                    run_cover_letter_agent)
                resume_path = "resume.pdf"
                cover_letter, error = run_cover_letter_agent(
                    resume_path, cover_jd, cover_company)

                if error:
                    st.error(f"❌ Error: {error}")
                else:
                    st.success(
                        "✅ Your cover letter is ready!")
                    st.subheader("👀 Cover Letter Preview")
                    st.markdown(cover_letter)
                    st.divider()

                    col1, col2 = st.columns(2)
                    with col1:
                        st.download_button(
                            label="⬇️ Download as TXT",
                            data=cover_letter,
                            file_name="cover_letter.txt",
                            mime="text/plain"
                        )
                    with col2:
                        if st.button(
                            "📄 Download as Word Doc",
                            key="cover_word"
                        ):
                            try:
                                from docx import Document
                                doc = Document()
                                for line in cover_letter.split('\n'):
                                    if not line.strip():
                                        doc.add_paragraph('')
                                        continue
                                    p = doc.add_paragraph()
                                    parts = re.split(
                                        r'\*\*(.*?)\*\*', line)
                                    for i, part in enumerate(parts):
                                        if i % 2 == 0:
                                            p.add_run(part)
                                        else:
                                            run = p.add_run(part)
                                            run.bold = True
                                doc_bytes = io.BytesIO()
                                doc.save(doc_bytes)
                                doc_bytes.seek(0)
                                st.download_button(
                                    label="⬇️ Save Word File",
                                    data=doc_bytes.getvalue(),
                                    file_name="cover_letter.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key="cover_docx"
                                )
                            except Exception as e:
                                st.error(f"❌ Error: {str(e)}")
            except Exception as e:
                st.error(
                    f"❌ Something went wrong: {str(e)}")
    else:
        st.warning(
            "⚠️ Please fill Company Name and Job Description!")

st.divider()

# ─────────────────────────────────────
# JOB SCOUT AGENT
# ─────────────────────────────────────
st.subheader("🔍 AI Smart Job Scout Agent")
st.markdown(
    "Searches 7 job variations, scores against "
    "your resume, shows matching keywords!")

# Load previous search preferences
try:
    from agents.scout_agent import load_search_preferences
    prefs = load_search_preferences()
    recent_roles = prefs.get("recent_roles", [])
    recent_locations = prefs.get("recent_locations", [])
except Exception:
    recent_roles = []
    recent_locations = []

# Show recent searches hint
if recent_roles:
    st.info(
        f"💡 Recent searches: "
        f"{', '.join(recent_roles[:3])}")

scout_col1, scout_col2 = st.columns(2)

with scout_col1:
    default_role = recent_roles[0] if recent_roles else ""
    job_role = st.text_input(
        "Job Role:",
        value=default_role,
        placeholder="e.g. Machine Learning Engineer",
        key="scout_role"
    )
    if recent_roles:
        selected_role = st.selectbox(
            "Or pick from recent searches:",
            [""] + recent_roles,
            key="role_select"
        )
        if selected_role:
            job_role = selected_role

with scout_col2:
    default_location = (
        recent_locations[0] if recent_locations else "")
    job_location = st.text_input(
        "Location:",
        value=default_location,
        placeholder="e.g. New York or Remote",
        key="scout_location"
    )
    if recent_locations:
        selected_location = st.selectbox(
            "Or pick from recent locations:",
            [""] + recent_locations,
            key="location_select"
        )
        if selected_location:
            job_location = selected_location

keywords_input = st.text_input(
    "Filter by keywords (optional, comma separated):",
    placeholder="e.g. Python, LangChain, Remote",
    key="scout_keywords"
)

show_new_only = st.checkbox(
    "Show only NEW jobs (hide previously seen)",
    value=False
)

if st.button("🔍 Find Best Matching Jobs"):
    if job_role and job_location:
        with st.spinner(
            "🔍 Searching 7 job variations and "
            "scoring against your resume... "
            "30-60 seconds..."
        ):
            try:
                from agents.scout_agent import (
                    run_scout_agent,
                    format_posted_time
                )

                keywords = None
                if keywords_input:
                    keywords = [
                        k.strip()
                        for k in keywords_input.split(",")
                    ]

                jobs, error = run_scout_agent(
                    job_role,
                    job_location,
                    keywords=keywords,
                    show_new_only=show_new_only,
                    pdf_path="resume.pdf"
                )

                if error:
                    st.error(f"❌ Error: {error}")
                elif not jobs:
                    st.warning(
                        "⚠️ No jobs found. "
                        "Try different search terms!")
                else:
                    new_count = len(
                        [j for j in jobs
                         if j.get("is_new", True)])
                    st.success(
                        f"✅ Found {len(jobs)} jobs "
                        f"({new_count} new)!")

                    for i, job in enumerate(jobs):
                        score = job.get("match_score", 0)
                        matching = job.get(
                            "matching_keywords", [])
                        missing = job.get(
                            "missing_keywords", [])

                        # Score labels
                        if score >= 80:
                            score_emoji = "🟢"
                            score_label = "Excellent Match"
                        elif score >= 65:
                            score_emoji = "🟡"
                            score_label = "Good Match"
                        else:
                            score_emoji = "🔴"
                            score_label = "Low Match"

                        is_new = job.get("is_new", True)
                        new_badge = "🆕 " if is_new else "👀 "
                        posted_fmt = format_posted_time(
                            job.get("posted", "N/A"))

                        with st.expander(
                            f"{new_badge}🏢 {job['company']} "
                            f"— {job['title']} | "
                            f"{score_emoji} {score}% "
                            f"{score_label}"
                        ):
                            col1, col2, col3 = st.columns(3)

                            with col1:
                                st.write(
                                    f"📍 **Location:** "
                                    f"{job['location']}")
                                st.write(
                                    f"⏰ **Posted:** "
                                    f"{posted_fmt}")

                            with col2:
                                st.write(
                                    f"💰 **Salary:** "
                                    f"{job['salary']}")
                                st.write(
                                    f"🎯 **Match:** {score}%")

                            with col3:
                                if job['apply_link'] != 'N/A':
                                    st.link_button(
                                        "🚀 Apply Now",
                                        job['apply_link']
                                    )

                            # Match score progress bar
                            st.progress(
                                min(score / 100, 1.0),
                                text=f"Resume Match: {score}%"
                            )

                            # Keyword Analysis Section
                            st.write(
                                "**🔍 Why this match score?**")

                            kw_col1, kw_col2 = st.columns(2)

                            with kw_col1:
                                if matching:
                                    st.success(
                                        "✅ **Keywords you have:**\n" +
                                        " • ".join(matching)
                                    )
                                else:
                                    st.warning(
                                        "⚠️ No matching keywords found")

                            with kw_col2:
                                if missing:
                                    st.error(
                                        "❌ **Missing keywords:**\n" +
                                        " • ".join(missing)
                                    )
                                    st.caption(
                                        "💡 Add these to your resume "
                                        "to improve match score!")
                                else:
                                    st.success(
                                        "🎉 No missing keywords!")

                            st.write(
                                "**📋 Job Description:**")
                            st.write(job['description'])

                            if st.button(
                                "➕ Add to Tracker",
                                key=f"scout_add_{i}"
                            ):
                                db = get_db()
                                new_app = JobApplication(
                                    company=job['company'],
                                    job_title=job['title'],
                                    job_url=job['apply_link'],
                                    location=job['location'],
                                    salary=job['salary'],
                                    status="Applied"
                                )
                                db.add(new_app)
                                db.commit()
                                db.close()
                                st.success(
                                    f"✅ {job['company']} "
                                    f"added to tracker!")

            except Exception as e:
                st.error(
                    f"❌ Something went wrong: {str(e)}")
    else:
        st.warning(
            "⚠️ Please fill in Job Role and Location!")