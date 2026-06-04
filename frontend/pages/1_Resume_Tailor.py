import sys
import os
import streamlit as st
import re
import io
import tempfile

sys.path.append(os.path.dirname(
    os.path.dirname(os.path.abspath(__file__))))

st.set_page_config(
    page_title="Resume Tailor",
    page_icon="🤖",
    layout="wide"
)

from frontend.theme import apply_dark_theme, handle_resume_upload
apply_dark_theme()

# ─────────────────────────────────────
# RESUME HANDLING
# Works on both laptop and Railway/cloud
# ─────────────────────────────────────

def get_resume_path():
    """
    Gets resume path from session state.
    Creates fresh temp file every time.
    Works on cloud where resume.pdf does not exist.
    """
    # Check session state first (uploaded via sidebar)
    if "resume_bytes" in st.session_state:
        with tempfile.NamedTemporaryFile(
            delete=False,
            suffix=".pdf"
        ) as tmp:
            tmp.write(st.session_state["resume_bytes"])
            return tmp.name

    # Fall back to local file (laptop only)
    if os.path.exists("resume.pdf"):
        return "resume.pdf"

    return None

# Handle resume upload in sidebar
resume_path = handle_resume_upload()
if resume_path:
    st.session_state["resume_path"] = resume_path

# ─────────────────────────────────────
# PAGE CONTENT
# ─────────────────────────────────────

st.title("🤖 AI Resume Tailor Agent")
st.markdown(
    "Paste a job description and AI will "
    "tailor your resume instantly!")

job_description = st.text_area(
    "Paste Job Description Here:",
    height=250,
    placeholder="Copy and paste the full job description here..."
)

if st.button("✨ Tailor My Resume with AI",
             use_container_width=True):
    if job_description:

        # Get fresh resume path every time
        resume_path = get_resume_path()

        if not resume_path:
            st.warning(
                "⚠️ Please upload your resume "
                "using the sidebar on the left!")
            st.stop()

        with st.spinner(
            "🤖 AI is tailoring your resume... "
            "30-60 seconds..."
        ):
            try:
                from agents.tailor_agent import run_tailor_agent

                tailored_resume, error = run_tailor_agent(
                    resume_path, job_description)

                if error:
                    st.error(f"❌ Error: {error}")
                else:
                    st.success("✅ Your tailored resume is ready!")
                    st.subheader("👀 Preview")
                    st.markdown(tailored_resume)
                    st.divider()

                    col1, col2 = st.columns(2)
                    with col1:
                        st.download_button(
                            label="⬇️ Download as TXT",
                            data=tailored_resume,
                            file_name="tailored_resume.txt",
                            mime="text/plain"
                        )
                    with col2:
                        if st.button("📄 Download as Word"):
                            try:
                                from docx import Document
                                doc = Document()
                                for line in tailored_resume.split('\n'):
                                    if not line.strip():
                                        doc.add_paragraph('')
                                        continue
                                    if (line.strip().startswith('*') and
                                            not line.strip().startswith('**')):
                                        p = doc.add_paragraph(
                                            style='List Bullet')
                                        line = line.strip().lstrip(
                                            '*').strip()
                                    else:
                                        p = doc.add_paragraph()
                                    parts = re.split(
                                        r'\*\*(.*?)\*\*', line)
                                    for i, part in enumerate(parts):
                                        if i % 2 == 0:
                                            p.add_run(part)
                                        else:
                                            run = p.add_run(part)
                                            run.bold = True
                                doc_bytes = io.BytesIO()
                                doc.save(doc_bytes)
                                doc_bytes.seek(0)
                                st.download_button(
                                    label="⬇️ Save Word File",
                                    data=doc_bytes.getvalue(),
                                    file_name="tailored_resume.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key="word_doc"
                                )
                            except Exception as e:
                                st.error(f"❌ {str(e)}")

            except Exception as e:
                st.error(f"❌ Something went wrong: {str(e)}")
    else:
        st.warning("⚠️ Please paste a job description first!")