import sys
import os
import streamlit as st
import re
import io
import tempfile

sys.path.append(os.path.dirname(
    os.path.dirname(os.path.abspath(__file__))))

st.set_page_config(
    page_title="Cover Letter",
    page_icon="✉️",
    layout="wide"
)

from frontend.theme import apply_dark_theme, handle_resume_upload
apply_dark_theme()

# ─────────────────────────────────────
# RESUME HANDLING
# Works on both laptop and Railway/cloud
# ─────────────────────────────────────

def get_resume_path():
    """
    Gets resume path from session state.
    Creates fresh temp file every time.
    Works on cloud where resume.pdf does not exist.
    """
    # Check session state first (uploaded via sidebar)
    if "resume_bytes" in st.session_state:
        with tempfile.NamedTemporaryFile(
            delete=False,
            suffix=".pdf"
        ) as tmp:
            tmp.write(st.session_state["resume_bytes"])
            return tmp.name

    # Fall back to local file (laptop only)
    if os.path.exists("resume.pdf"):
        return "resume.pdf"

    return None

# Handle resume upload in sidebar
resume_path = handle_resume_upload()
if resume_path:
    st.session_state["resume_path"] = resume_path

# ─────────────────────────────────────
# PAGE CONTENT
# ─────────────────────────────────────

st.title("✉️ AI Cover Letter Agent")
st.markdown(
    "Enter company name and job description — "
    "AI writes your personalized cover letter!")

col1, col2 = st.columns(2)

with col1:
    company = st.text_input(
        "Company Name:",
        placeholder="e.g. Google, Amazon, Microsoft"
    )

with col2:
    st.write("")
    st.write("")
    st.info("💡 AI will personalize the letter for this company!")

job_description = st.text_area(
    "Paste Job Description Here:",
    height=250,
    placeholder="Copy and paste the full job description here..."
)

if st.button("✉️ Write My Cover Letter",
             use_container_width=True):
    if company and job_description:

        # Get fresh resume path every time
        resume_path = get_resume_path()

        if not resume_path:
            st.warning(
                "⚠️ Please upload your resume "
                "using the sidebar on the left!")
            st.stop()

        with st.spinner(
            "✍️ AI is writing your cover letter... "
            "10-30 seconds..."
        ):
            try:
                from agents.cover_letter_agent import run_cover_letter_agent

                cover_letter, error = run_cover_letter_agent(
                    resume_path, job_description, company)

                if error:
                    st.error(f"❌ Error: {error}")
                else:
                    st.success("✅ Your cover letter is ready!")
                    st.subheader("👀 Preview")
                    st.markdown(cover_letter)
                    st.divider()

                    col1, col2 = st.columns(2)
                    with col1:
                        st.download_button(
                            label="⬇️ Download as TXT",
                            data=cover_letter,
                            file_name="cover_letter.txt",
                            mime="text/plain"
                        )
                    with col2:
                        if st.button("📄 Download as Word"):
                            try:
                                from docx import Document
                                doc = Document()
                                for line in cover_letter.split('\n'):
                                    if not line.strip():
                                        doc.add_paragraph('')
                                        continue
                                    p = doc.add_paragraph()
                                    parts = re.split(
                                        r'\*\*(.*?)\*\*', line)
                                    for i, part in enumerate(parts):
                                        if i % 2 == 0:
                                            p.add_run(part)
                                        else:
                                            run = p.add_run(part)
                                            run.bold = True
                                doc_bytes = io.BytesIO()
                                doc.save(doc_bytes)
                                doc_bytes.seek(0)
                                st.download_button(
                                    label="⬇️ Save Word File",
                                    data=doc_bytes.getvalue(),
                                    file_name="cover_letter.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key="cover_word"
                                )
                            except Exception as e:
                                st.error(f"❌ {str(e)}")

            except Exception as e:
                st.error(f"❌ Something went wrong: {str(e)}")
    else:
        st.warning(
            "⚠️ Please fill Company Name and Job Description!")