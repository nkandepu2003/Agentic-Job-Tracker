# pages/6_Smart_Apply.py
import streamlit as st
import sys
import os
import io
import re
import json

sys.path.append(os.path.dirname(
    os.path.dirname(os.path.abspath(__file__))))

st.set_page_config(
    page_title="Smart Apply",
    page_icon="🚀",
    layout="wide"
)

from frontend.theme import apply_dark_theme
apply_dark_theme()

# ─────────────────────────────────────
# RECENT SEARCH FUNCTIONS
# ─────────────────────────────────────

SMART_APPLY_PREFS = "smart_apply_preferences.json"


def save_smart_apply_prefs(company, role):
    try:
        if os.path.exists(SMART_APPLY_PREFS):
            with open(SMART_APPLY_PREFS, "r") as f:
                prefs = json.load(f)
        else:
            prefs = {"companies": [], "roles": []}
        if company and company not in prefs["companies"]:
            prefs["companies"].insert(0, company)
            prefs["companies"] = prefs["companies"][:5]
        if role and role not in prefs["roles"]:
            prefs["roles"].insert(0, role)
            prefs["roles"] = prefs["roles"][:5]
        with open(SMART_APPLY_PREFS, "w") as f:
            json.dump(prefs, f)
    except Exception as e:
        print(f"Could not save prefs: {str(e)}")


def load_smart_apply_prefs():
    try:
        if os.path.exists(SMART_APPLY_PREFS):
            with open(SMART_APPLY_PREFS, "r") as f:
                return json.load(f)
        return {"companies": [], "roles": []}
    except Exception:
        return {"companies": [], "roles": []}


# ─────────────────────────────────────
# PAGE TITLE
# ─────────────────────────────────────

st.title("🚀 Smart Apply — Powered by LangGraph")
st.markdown(
    "One click → AI scores your match, "
    "tailors resume, writes cover letter, "
    "and logs application automatically!")

st.info(
    "💡 This uses LangGraph to connect all "
    "AI agents in an automatic workflow. "
    "Just fill in the details and click apply!")

# ─────────────────────────────────────
# LOAD RECENT SEARCHES
# ─────────────────────────────────────

sa_prefs = load_smart_apply_prefs()
sa_companies = sa_prefs.get("companies", [])
sa_roles = sa_prefs.get("roles", [])

# ─────────────────────────────────────
# INPUT FIELDS
# ─────────────────────────────────────

# Recent company selection
if sa_companies:
    st.markdown(
        f'<p style="color:#6b6783; font-size:12px;">💡 Recent: {", ".join(sa_companies[:3])}</p>',
        unsafe_allow_html=True
    )

col1, col2 = st.columns(2)

with col1:
    if sa_companies:
        company_choice = st.selectbox(
            "Company Name:",
            ["Type a new company..."] + sa_companies,
            key="sa_company_select"
        )
        if company_choice == "Type a new company...":
            company_name = st.text_input(
                "Enter company name:",
                placeholder="e.g. Google",
                key="sa_company_input"
            )
        else:
            company_name = company_choice
    else:
        company_name = st.text_input(
            "Company Name:",
            placeholder="e.g. Google",
            key="sa_company_input"
        )

with col2:
    if sa_roles:
        role_choice = st.selectbox(
            "Job Title:",
            ["Type a new role..."] + sa_roles,
            key="sa_role_select"
        )
        if role_choice == "Type a new role...":
            job_title = st.text_input(
                "Enter job title:",
                placeholder="e.g. ML Engineer",
                key="sa_role_input"
            )
        else:
            job_title = role_choice
    else:
        job_title = st.text_input(
            "Job Title:",
            placeholder="e.g. ML Engineer",
            key="sa_role_input"
        )

job_url = st.text_input(
    "Job URL (optional):",
    placeholder="https://careers.google.com/...",
    key="sa_url_input"
)

job_description = st.text_area(
    "Paste Full Job Description:",
    height=250,
    placeholder="Copy and paste the complete job description here...",
    key="sa_jd_input"
)

st.divider()

# ─────────────────────────────────────
# SMART APPLY BUTTON
# ─────────────────────────────────────

if st.button(
    "🚀 Smart Apply — Run All Agents",
    use_container_width=True,
    type="primary"
):
    if company_name and job_title and job_description:

        progress_bar = st.progress(0)
        status_text = st.empty()

        with st.spinner(
            "🤖 LangGraph is running all agents... "
            "This takes 1-2 minutes..."
        ):
            try:
                from agents.job_application_graph import (
                    run_application_graph)

                status_text.text(
                    "Starting agentic workflow...")
                progress_bar.progress(10)

                final_state, error = run_application_graph(
                    job_description=job_description,
                    company_name=company_name,
                    job_title=job_title,
                    job_url=job_url or "N/A",
                    resume_path="resume.pdf"
                )

                save_smart_apply_prefs(
                    company_name, job_title)

                progress_bar.progress(100)
                status_text.text("All agents complete!")

                if error:
                    st.error(f"❌ Error: {error}")
                else:
                    st.success(
                        "🎉 Smart Apply Complete! "
                        "Everything is ready!")

                    tab1, tab2, tab3, tab4 = st.tabs([
                        "🎯 Match Score",
                        "📄 Tailored Resume",
                        "✉️ Cover Letter",
                        "📊 Application Log"
                    ])

                    with tab1:
                        st.subheader("🎯 Job Match Analysis")
                        score = final_state.get(
                            "match_score", 0)

                        if score >= 70:
                            st.success(
                                f"🟢 Excellent Match: {score}%")
                        elif score >= 50:
                            st.warning(
                                f"🟡 Good Match: {score}%")
                        else:
                            st.error(
                                f"🔴 Low Match: {score}%")

                        st.progress(
                            min(score / 100, 1.0),
                            text=f"Resume Match: {score}%"
                        )

                        kw_col1, kw_col2 = st.columns(2)
                        with kw_col1:
                            matching = final_state.get(
                                "matching_keywords", [])
                            if matching:
                                st.success(
                                    "✅ Keywords you have:\n" +
                                    " • ".join(matching))
                        with kw_col2:
                            missing = final_state.get(
                                "missing_keywords", [])
                            if missing:
                                st.error(
                                    "❌ Missing keywords:\n" +
                                    " • ".join(missing))
                                st.caption(
                                    "💡 Add these to your "
                                    "resume to improve score!")

                    with tab2:
                        st.subheader("📄 Tailored Resume")
                        tailored = final_state.get(
                            "tailored_resume")

                        if tailored:
                            st.markdown(tailored)
                            st.divider()
                            col1, col2 = st.columns(2)
                            with col1:
                                st.download_button(
                                    label="⬇️ Download as TXT",
                                    data=tailored,
                                    file_name=f"{company_name}_resume.txt",
                                    mime="text/plain",
                                    key="resume_txt"
                                )
                            with col2:
                                try:
                                    from docx import Document
                                    doc = Document()
                                    for line in tailored.split('\n'):
                                        if not line.strip():
                                            doc.add_paragraph('')
                                            continue
                                        if (line.strip().startswith('*')
                                                and not line.strip().startswith('**')):
                                            p = doc.add_paragraph(
                                                style='List Bullet')
                                            line = line.strip().lstrip(
                                                '*').strip()
                                        else:
                                            p = doc.add_paragraph()
                                        parts = re.split(
                                            r'\*\*(.*?)\*\*', line)
                                        for i, part in enumerate(parts):
                                            if i % 2 == 0:
                                                p.add_run(part)
                                            else:
                                                run = p.add_run(part)
                                                run.bold = True
                                    doc_bytes = io.BytesIO()
                                    doc.save(doc_bytes)
                                    doc_bytes.seek(0)
                                    st.download_button(
                                        label="⬇️ Download as Word",
                                        data=doc_bytes.getvalue(),
                                        file_name=f"{company_name}_resume.docx",
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                        key="resume_word"
                                    )
                                except Exception as e:
                                    st.error(f"Word error: {str(e)}")
                        else:
                            st.warning(
                                "⚠️ Resume tailoring failed.")

                    with tab3:
                        st.subheader("✉️ Cover Letter")
                        cover = final_state.get("cover_letter")
                        if cover:
                            st.markdown(cover)
                            st.divider()
                            st.download_button(
                                label="⬇️ Download Cover Letter",
                                data=cover,
                                file_name=f"{company_name}_cover_letter.txt",
                                mime="text/plain",
                                key="cover_dl"
                            )
                        else:
                            st.warning(
                                "⚠️ Cover letter generation failed.")

                    with tab4:
                        st.subheader("📊 Application Status")
                        logged = final_state.get(
                            "application_logged", False)
                        if logged:
                            st.success(
                                "✅ Application automatically "
                                "logged to your tracker!")
                            st.write(
                                f"**Company:** {company_name}")
                            st.write(f"**Role:** {job_title}")
                            st.write("**Status:** Applied")
                            st.write(
                                f"**Match Score:** "
                                f"{final_state.get('match_score', 0)}%")
                            st.info(
                                "👉 Go to Dashboard to "
                                "see your application!")
                        else:
                            st.warning(
                                "⚠️ Application logging failed. "
                                "Add manually from Dashboard.")

            except Exception as e:
                st.error(
                    f"❌ Something went wrong: {str(e)}")
                progress_bar.progress(0)

    else:
        st.warning(
            "⚠️ Please fill Company Name, "
            "Job Title and Job Description!")